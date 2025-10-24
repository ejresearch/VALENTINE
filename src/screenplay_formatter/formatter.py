"""Formatter modules for different output formats."""

import os
from abc import ABC, abstractmethod
from typing import List, Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from .parser import ScreenplayElement, ElementType


class BaseFormatter(ABC):
    """Base class for screenplay formatters."""

    # Standard screenplay formatting measurements
    PAGE_WIDTH = 8.5  # inches
    PAGE_HEIGHT = 11   # inches
    LEFT_MARGIN = 1.5  # inches
    RIGHT_MARGIN = 1.0  # inches
    TOP_MARGIN = 1.0   # inches
    BOTTOM_MARGIN = 1.0  # inches

    # Element indentations (from left margin)
    CHARACTER_INDENT = 3.7  # inches from left edge
    DIALOGUE_LEFT_INDENT = 2.5  # inches from left edge
    DIALOGUE_RIGHT_INDENT = 1.5  # inches from right edge
    PARENTHETICAL_INDENT = 3.1  # inches from left edge
    TRANSITION_INDENT = 6.0  # inches from left edge (flush right)

    # Text formatting
    FONT_NAME = "Courier"
    FONT_SIZE = 12

    @abstractmethod
    def format(self, elements: List[ScreenplayElement], output_path: str):
        """Format screenplay elements and save to file."""
        pass


class TextFormatter(BaseFormatter):
    """Format screenplay as plain text with spacing."""

    def __init__(self, include_scene_numbers: bool = False):
        # Calculate character positions for 80-character width
        self.page_width_chars = 80
        self.left_margin_chars = 15
        self.character_position = 37  # Centered around column 37
        self.dialogue_left_indent = 25
        self.dialogue_right_margin = 65
        self.parenthetical_indent = 31
        self.transition_position = 70
        self.include_scene_numbers = include_scene_numbers

    def format(self, elements: List[ScreenplayElement], output_path: str):
        """Format screenplay elements as plain text."""
        lines = []

        # Separate title page elements from screenplay body
        title_page_elements = [e for e in elements if e.type in [
            ElementType.TITLE_PAGE_TITLE, ElementType.TITLE_PAGE_AUTHOR,
            ElementType.TITLE_PAGE_CONTACT, ElementType.TITLE_PAGE_CREDIT
        ]]
        screenplay_elements = [e for e in elements if e not in title_page_elements]

        # Format title page if present
        if title_page_elements:
            lines.extend(self._format_title_page(title_page_elements))
            lines.append("")  # Page break after title page

        for i, element in enumerate(screenplay_elements):
            formatted = self._format_element(element)
            if formatted:
                lines.extend(formatted)

            # Add spacing between elements
            if self._needs_spacing_after(element, screenplay_elements, i):
                lines.append("")

        # Write to file
        with open(output_path, 'w') as f:
            f.write('\n'.join(lines))

    def _format_title_page(self, elements: List[ScreenplayElement]) -> List[str]:
        """Format title page elements."""
        lines = []

        # Extract elements
        title = None
        author = None
        contact = []
        credit = None

        for elem in elements:
            if elem.type == ElementType.TITLE_PAGE_TITLE:
                title = elem.content
            elif elem.type == ElementType.TITLE_PAGE_AUTHOR:
                author = elem.content
            elif elem.type == ElementType.TITLE_PAGE_CONTACT:
                contact.append(elem.content)
            elif elem.type == ElementType.TITLE_PAGE_CREDIT:
                credit = elem.content

        # Format title page (industry standard layout)
        # Proper order: Title, then credit, then author
        lines.extend([""] * 10)  # Vertical spacing from top

        if title:
            lines.append(title.upper().center(self.page_width_chars).rstrip())
            lines.append("")
            lines.append("")
            lines.append("")

        if credit:
            lines.append(credit.center(self.page_width_chars).rstrip())
            lines.append("")

        if author:
            lines.append(author.center(self.page_width_chars).rstrip())

        # Add more spacing to push contact info to bottom
        lines.extend([""] * 15)

        # Contact info (bottom right)
        if contact:
            for contact_line in contact:
                lines.append(contact_line.rjust(self.page_width_chars - 5))

        return lines

    def _format_element(self, element: ScreenplayElement) -> List[str]:
        """Format a single screenplay element."""
        if element.type == ElementType.BLANK:
            return []

        if element.type == ElementType.SCENE_HEADING:
            scene_heading = element.content.upper()
            if self.include_scene_numbers and element.scene_number:
                # Add scene number on both sides (industry standard)
                scene_heading = f"{element.scene_number}   {scene_heading}   {element.scene_number}"
            return [scene_heading]

        if element.type == ElementType.ACTION:
            return self._wrap_text(element.content, 0, self.page_width_chars)

        if element.type == ElementType.CHARACTER:
            # Center the character name
            return [element.content.upper().center(self.page_width_chars).rstrip()]

        if element.type == ElementType.DIALOGUE:
            return self._wrap_text(
                element.content,
                self.dialogue_left_indent,
                self.dialogue_right_margin
            )

        if element.type == ElementType.PARENTHETICAL:
            return self._wrap_text(
                element.content,
                self.parenthetical_indent,
                self.dialogue_right_margin
            )

        if element.type == ElementType.TRANSITION:
            # FADE IN: is left-aligned, all other transitions are right-aligned
            if element.content.upper().strip() == "FADE IN:":
                return [element.content.upper()]
            else:
                # Right-align other transitions
                return [element.content.upper().rjust(self.transition_position)]

        if element.type in [ElementType.MONTAGE_BEGIN, ElementType.MONTAGE_END]:
            return [element.content.upper()]

        if element.type in [ElementType.TITLE, ElementType.CHYRON]:
            return [element.content.upper()]

        if element.type == ElementType.SHOT:
            return [element.content.upper()]

        if element.type == ElementType.PAGE_BREAK:
            return ["\n" * 3]  # Force page break with extra spacing

        if element.type in [ElementType.DUAL_DIALOGUE_LEFT, ElementType.DUAL_DIALOGUE_RIGHT]:
            # Dual dialogue - these will be handled specially in post-processing
            return [element.content.upper().center(self.page_width_chars // 2).rstrip()]

        if element.type == ElementType.VFX_SFX:
            # Format VFX/SFX as action but keep brackets
            return [element.content.upper()]

        if element.type == ElementType.MORE:
            # Format (MORE) centered
            return [element.content.center(self.page_width_chars).rstrip()]

        return [element.content]

    def _wrap_text(self, text: str, left_indent: int, right_margin: int) -> List[str]:
        """Wrap text to fit within margins."""
        width = right_margin - left_indent
        words = text.split()
        lines = []
        current_line = []
        current_length = 0

        for word in words:
            if current_length + len(word) + 1 > width:
                if current_line:
                    lines.append(' ' * left_indent + ' '.join(current_line))
                current_line = [word]
                current_length = len(word)
            else:
                current_line.append(word)
                current_length += len(word) + 1

        if current_line:
            lines.append(' ' * left_indent + ' '.join(current_line))

        return lines

    def _needs_spacing_after(self, element: ScreenplayElement,
                            elements: List[ScreenplayElement],
                            index: int) -> bool:
        """Determine if spacing is needed after an element."""
        if element.type == ElementType.BLANK:
            return False

        # Look ahead to next non-blank element
        next_element = None
        for i in range(index + 1, len(elements)):
            if elements[i].type != ElementType.BLANK:
                next_element = elements[i]
                break

        if not next_element:
            return False

        # Add spacing after scene headings
        if element.type == ElementType.SCENE_HEADING:
            return True

        # Add spacing after shot headers
        if element.type == ElementType.SHOT:
            return True

        # Add spacing after action blocks
        if element.type == ElementType.ACTION and next_element.type != ElementType.ACTION:
            return True

        # Add spacing after dialogue blocks
        if element.type in [ElementType.DIALOGUE, ElementType.PARENTHETICAL]:
            if next_element.type not in [ElementType.DIALOGUE, ElementType.PARENTHETICAL]:
                return True

        # Add spacing after transitions
        if element.type == ElementType.TRANSITION:
            return True

        # Add spacing before CHARACTER names (professional standard)
        if next_element.type in [ElementType.CHARACTER, ElementType.DUAL_DIALOGUE_LEFT, ElementType.DUAL_DIALOGUE_RIGHT]:
            if element.type not in [ElementType.PARENTHETICAL, ElementType.DIALOGUE]:
                return True

        # No spacing after page breaks (they create their own)
        if element.type == ElementType.PAGE_BREAK:
            return False

        return False


class DocxFormatter(BaseFormatter):
    """Format screenplay as DOCX with proper styles."""

    def __init__(self, include_scene_numbers: bool = False):
        super().__init__()
        self.include_scene_numbers = include_scene_numbers

    def format(self, elements: List[ScreenplayElement], output_path: str):
        """Format screenplay elements as DOCX."""
        doc = Document()
        self._setup_page(doc)
        self._create_styles(doc)

        # Separate title page elements from screenplay body
        title_page_elements = [e for e in elements if e.type in [
            ElementType.TITLE_PAGE_TITLE, ElementType.TITLE_PAGE_AUTHOR,
            ElementType.TITLE_PAGE_CONTACT, ElementType.TITLE_PAGE_CREDIT
        ]]
        screenplay_elements = [e for e in elements if e not in title_page_elements]

        # Add title page if present
        if title_page_elements:
            self._add_title_page(doc, title_page_elements)
            doc.add_page_break()

        for i, element in enumerate(screenplay_elements):
            self._add_element(doc, element)

            # Add spacing between elements
            if self._needs_spacing_after(element, screenplay_elements, i):
                doc.add_paragraph()

        doc.save(output_path)

    def _setup_page(self, doc: Document):
        """Set up page layout with page numbering."""
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(self.PAGE_HEIGHT)
            section.page_width = Inches(self.PAGE_WIDTH)
            section.left_margin = Inches(self.LEFT_MARGIN)
            section.right_margin = Inches(self.RIGHT_MARGIN)
            section.top_margin = Inches(self.TOP_MARGIN)
            section.bottom_margin = Inches(self.BOTTOM_MARGIN)

            # Add page numbering in header (top right, industry standard)
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header_run = header_para.add_run()

            # Add page number field (shows current page number)
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            header_run._r.append(fldChar1)
            header_run._r.append(instrText)
            header_run._r.append(fldChar2)
            header_run.font.name = self.FONT_NAME
            header_run.font.size = Pt(self.FONT_SIZE)
            header_para.add_run('.')  # Add period after page number

    def _add_title_page(self, doc: Document, elements: List[ScreenplayElement]):
        """Add title page to document."""
        # Extract elements
        title = None
        author = None
        contact = []
        credit = None

        for elem in elements:
            if elem.type == ElementType.TITLE_PAGE_TITLE:
                title = elem.content
            elif elem.type == ElementType.TITLE_PAGE_AUTHOR:
                author = elem.content
            elif elem.type == ElementType.TITLE_PAGE_CONTACT:
                contact.append(elem.content)
            elif elem.type == ElementType.TITLE_PAGE_CREDIT:
                credit = elem.content

        # Add vertical spacing
        for _ in range(10):
            doc.add_paragraph()

        # Title (centered, uppercase)
        if title:
            title_para = doc.add_paragraph(title.upper())
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.runs[0].font.name = self.FONT_NAME
            title_para.runs[0].font.size = Pt(self.FONT_SIZE)
            title_para.runs[0].font.bold = False
            doc.add_paragraph()

        # Credit line (centered)
        if credit:
            credit_para = doc.add_paragraph(credit)
            credit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            credit_para.runs[0].font.name = self.FONT_NAME
            credit_para.runs[0].font.size = Pt(self.FONT_SIZE)
            doc.add_paragraph()

        # Author (centered)
        if author:
            author_para = doc.add_paragraph(author)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.runs[0].font.name = self.FONT_NAME
            author_para.runs[0].font.size = Pt(self.FONT_SIZE)

        # Add more spacing to push contact info down
        for _ in range(15):
            doc.add_paragraph()

        # Contact info (bottom right)
        if contact:
            for contact_line in contact:
                contact_para = doc.add_paragraph(contact_line)
                contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                contact_para.runs[0].font.name = self.FONT_NAME
                contact_para.runs[0].font.size = Pt(self.FONT_SIZE)

    def _create_styles(self, doc: Document):
        """Create custom styles for screenplay elements."""
        styles = doc.styles

        # Scene Heading style
        scene_style = styles.add_style('SceneHeading', WD_STYLE_TYPE.PARAGRAPH)
        scene_style.font.name = self.FONT_NAME
        scene_style.font.size = Pt(self.FONT_SIZE)
        scene_style.font.bold = False
        scene_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        scene_style.paragraph_format.space_after = Pt(12)

        # Action style
        action_style = styles.add_style('Action', WD_STYLE_TYPE.PARAGRAPH)
        action_style.font.name = self.FONT_NAME
        action_style.font.size = Pt(self.FONT_SIZE)
        action_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        action_style.paragraph_format.space_after = Pt(12)

        # Character style
        char_style = styles.add_style('Character', WD_STYLE_TYPE.PARAGRAPH)
        char_style.font.name = self.FONT_NAME
        char_style.font.size = Pt(self.FONT_SIZE)
        char_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        char_style.paragraph_format.space_before = Pt(12)
        char_style.paragraph_format.space_after = Pt(0)

        # Dialogue style
        dialogue_style = styles.add_style('Dialogue', WD_STYLE_TYPE.PARAGRAPH)
        dialogue_style.font.name = self.FONT_NAME
        dialogue_style.font.size = Pt(self.FONT_SIZE)
        dialogue_style.paragraph_format.left_indent = Inches(1.0)
        dialogue_style.paragraph_format.right_indent = Inches(1.5)
        dialogue_style.paragraph_format.space_after = Pt(0)

        # Parenthetical style
        paren_style = styles.add_style('Parenthetical', WD_STYLE_TYPE.PARAGRAPH)
        paren_style.font.name = self.FONT_NAME
        paren_style.font.size = Pt(self.FONT_SIZE)
        paren_style.paragraph_format.left_indent = Inches(1.6)
        paren_style.paragraph_format.right_indent = Inches(2.0)
        paren_style.paragraph_format.space_after = Pt(0)

        # Transition style
        trans_style = styles.add_style('Transition', WD_STYLE_TYPE.PARAGRAPH)
        trans_style.font.name = self.FONT_NAME
        trans_style.font.size = Pt(self.FONT_SIZE)
        trans_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        trans_style.paragraph_format.space_before = Pt(12)
        trans_style.paragraph_format.space_after = Pt(12)

        # Shot style (new)
        shot_style = styles.add_style('Shot', WD_STYLE_TYPE.PARAGRAPH)
        shot_style.font.name = self.FONT_NAME
        shot_style.font.size = Pt(self.FONT_SIZE)
        shot_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shot_style.paragraph_format.space_before = Pt(12)
        shot_style.paragraph_format.space_after = Pt(12)

    def _add_element(self, doc: Document, element: ScreenplayElement):
        """Add a screenplay element to the document."""
        if element.type == ElementType.BLANK:
            return

        # Handle page breaks
        if element.type == ElementType.PAGE_BREAK:
            doc.add_page_break()
            return

        style_map = {
            ElementType.SCENE_HEADING: 'SceneHeading',
            ElementType.ACTION: 'Action',
            ElementType.CHARACTER: 'Character',
            ElementType.DIALOGUE: 'Dialogue',
            ElementType.PARENTHETICAL: 'Parenthetical',
            ElementType.TRANSITION: 'Transition',
            ElementType.SHOT: 'Shot',
            ElementType.DUAL_DIALOGUE_LEFT: 'Character',
            ElementType.DUAL_DIALOGUE_RIGHT: 'Character',
            ElementType.MONTAGE_BEGIN: 'SceneHeading',
            ElementType.MONTAGE_END: 'SceneHeading',
            ElementType.TITLE: 'Action',
            ElementType.CHYRON: 'Action',
            ElementType.VFX_SFX: 'Action',
            ElementType.MORE: 'Character',
        }

        style = style_map.get(element.type, 'Normal')

        # Format content based on type
        content = element.content
        if element.type in [ElementType.SCENE_HEADING, ElementType.CHARACTER,
                           ElementType.TRANSITION, ElementType.MONTAGE_BEGIN,
                           ElementType.MONTAGE_END, ElementType.VFX_SFX]:
            content = content.upper()

        # Add scene numbers if enabled
        if element.type == ElementType.SCENE_HEADING and self.include_scene_numbers and element.scene_number:
            content = f"{element.scene_number}   {content}   {element.scene_number}"

        paragraph = doc.add_paragraph(content, style=style)

        # Add page break protection for dialogue blocks and scene headings
        if element.type == ElementType.CHARACTER:
            # Keep character name with following dialogue
            paragraph.paragraph_format.keep_with_next = True
        elif element.type == ElementType.DIALOGUE:
            # Keep dialogue lines together
            paragraph.paragraph_format.keep_with_next = True
        elif element.type == ElementType.PARENTHETICAL:
            # Keep parentheticals with character/dialogue
            paragraph.paragraph_format.keep_with_next = True
        elif element.type == ElementType.SCENE_HEADING:
            # Keep scene headings with following action
            paragraph.paragraph_format.keep_with_next = True

    def _needs_spacing_after(self, element: ScreenplayElement,
                            elements: List[ScreenplayElement],
                            index: int) -> bool:
        """Determine if spacing is needed after an element."""
        # Similar logic to TextFormatter but adjusted for DOCX
        return False  # Spacing is handled by styles in DOCX


class PdfFormatter(BaseFormatter):
    """Format screenplay as PDF with pagination."""

    def __init__(self, include_scene_numbers: bool = False):
        super().__init__()
        self.lines_per_page = 55
        self.line_height = 12  # points
        self.current_page = 1
        self.current_y = 0
        self.include_scene_numbers = include_scene_numbers

    def format(self, elements: List[ScreenplayElement], output_path: str):
        """Format screenplay elements as PDF."""
        c = canvas.Canvas(output_path, pagesize=letter)

        # Try to register Courier font
        try:
            from reportlab.pdfbase.pdfmetrics import registerFont
            from reportlab.pdfbase.ttfonts import TTFont
            # This would need the actual Courier font file
            c.setFont("Courier", self.FONT_SIZE)
        except:
            # Fall back to Courier which should be built-in
            c.setFont("Courier", self.FONT_SIZE)

        self._start_new_page(c)

        for i, element in enumerate(elements):
            self._add_element(c, element, elements, i)

        c.save()

    def _start_new_page(self, c: canvas.Canvas):
        """Start a new page."""
        if self.current_page > 1:
            c.showPage()
            c.setFont("Courier", self.FONT_SIZE)
            # Add page number
            c.drawRightString(
                self.PAGE_WIDTH * inch - self.RIGHT_MARGIN * inch,
                self.PAGE_HEIGHT * inch - 0.5 * inch,
                str(self.current_page)
            )

        self.current_y = self.PAGE_HEIGHT * inch - self.TOP_MARGIN * inch
        self.current_page += 1

    def _add_element(self, c: canvas.Canvas, element: ScreenplayElement,
                    elements: List[ScreenplayElement], index: int):
        """Add element to PDF."""
        if element.type == ElementType.BLANK:
            self._move_down(c, 1)
            return

        # Check for blocks that shouldn't break across pages
        if element.type == ElementType.CHARACTER:
            dialogue_block_height = self._calculate_dialogue_block_height(elements, index)
            if dialogue_block_height > 0 and self.current_y - dialogue_block_height < self.BOTTOM_MARGIN * inch:
                # Start new page if dialogue block won't fit
                self._start_new_page(c)
        elif element.type == ElementType.SCENE_HEADING:
            # Scene headings need at least 3 lines of space after them
            required_space = self.line_height * 3
            if self.current_y - required_space < self.BOTTOM_MARGIN * inch:
                self._start_new_page(c)

        # Format based on element type
        if element.type == ElementType.SCENE_HEADING:
            scene_heading = element.content.upper()
            if self.include_scene_numbers and element.scene_number:
                scene_heading = f"{element.scene_number}   {scene_heading}   {element.scene_number}"
            self._add_text(c, scene_heading, self.LEFT_MARGIN * inch)
            self._move_down(c, 2)

        elif element.type == ElementType.ACTION:
            lines = self._wrap_pdf_text(element.content, 65)
            for line in lines:
                self._add_text(c, line, self.LEFT_MARGIN * inch)
                self._move_down(c, 1)
            self._move_down(c, 1)

        elif element.type == ElementType.CHARACTER:
            # Center character name
            text_width = c.stringWidth(element.content.upper(), "Courier", self.FONT_SIZE)
            x_pos = (self.PAGE_WIDTH * inch) / 2 - text_width / 2
            self._add_text(c, element.content.upper(), x_pos)
            self._move_down(c, 1)

        elif element.type == ElementType.DIALOGUE:
            lines = self._wrap_pdf_text(element.content, 35)
            for line in lines:
                self._add_text(c, line, self.DIALOGUE_LEFT_INDENT * inch)
                self._move_down(c, 1)

        elif element.type == ElementType.PARENTHETICAL:
            lines = self._wrap_pdf_text(element.content, 25)
            for line in lines:
                self._add_text(c, line, self.PARENTHETICAL_INDENT * inch)
                self._move_down(c, 1)

        elif element.type == ElementType.TRANSITION:
            # FADE IN: is left-aligned, all others are right-aligned
            if element.content.upper().strip() == "FADE IN:":
                self._add_text(c, element.content.upper(), self.LEFT_MARGIN * inch)
            else:
                # Right align other transitions
                text_width = c.stringWidth(element.content.upper(), "Courier", self.FONT_SIZE)
                x_pos = self.PAGE_WIDTH * inch - self.RIGHT_MARGIN * inch - text_width
                self._add_text(c, element.content.upper(), x_pos)
            self._move_down(c, 2)

    def _add_text(self, c: canvas.Canvas, text: str, x_pos: float):
        """Add text at current y position."""
        c.drawString(x_pos, self.current_y, text)

    def _move_down(self, c: canvas.Canvas, lines: int):
        """Move down by specified number of lines."""
        self.current_y -= self.line_height * lines

        # Check if we need a new page
        if self.current_y < self.BOTTOM_MARGIN * inch:
            self._start_new_page(c)

    def _wrap_pdf_text(self, text: str, max_chars: int) -> List[str]:
        """Wrap text to fit within character limit."""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0

        for word in words:
            if current_length + len(word) + 1 > max_chars:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
            else:
                current_line.append(word)
                current_length += len(word) + 1

        if current_line:
            lines.append(' '.join(current_line))

        return lines

    def _calculate_dialogue_block_height(self, elements: List[ScreenplayElement], start_index: int) -> float:
        """Calculate the height needed for a complete dialogue block."""
        if start_index >= len(elements) or elements[start_index].type != ElementType.CHARACTER:
            return 0

        height = 0
        i = start_index

        # Count lines for character name
        height += self.line_height

        # Look ahead for dialogue, parentheticals, and continuation
        while i + 1 < len(elements):
            next_elem = elements[i + 1]

            if next_elem.type == ElementType.DIALOGUE:
                # Count wrapped lines for dialogue
                lines = self._wrap_pdf_text(next_elem.content, 35)
                height += len(lines) * self.line_height
            elif next_elem.type == ElementType.PARENTHETICAL:
                # Count wrapped lines for parenthetical
                lines = self._wrap_pdf_text(next_elem.content, 25)
                height += len(lines) * self.line_height
            elif next_elem.type == ElementType.BLANK:
                # Small space between dialogue elements
                height += self.line_height * 0.5
            elif next_elem.type == ElementType.CHARACTER and "(CONT'D)" in next_elem.content:
                # Character continuation - part of same conversation
                height += self.line_height
            else:
                # End of dialogue block
                break

            i += 1

        # Add minimum buffer space
        height += self.line_height * 2

        return height